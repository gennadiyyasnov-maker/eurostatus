import docx
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

def add_row(table, text_bg, text_ru, bold=False, align_center=False):
    cells = table.add_row().cells
    
    p0 = cells[0].paragraphs[0]
    p1 = cells[1].paragraphs[0]
    
    p0.add_run(text_bg).bold = bold
    p1.add_run(text_ru).bold = bold
    
    if align_center:
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p0.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

def create_procura_columns():
    doc = docx.Document()

    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    # Set styling
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)

    # Add a 2-column table
    table = doc.add_table(rows=0, cols=2)
    table.autofit = False
    for col in table.columns:
        col.width = Cm(8.5)

    # Invisible borders logic (by default python-docx tables have no borders, which is what we want for a clean look)

    # 1. Title
    add_row(table, 'ПЪЛНОМОЩНО', 'ДОВЕРЕННОСТЬ', bold=True, align_center=True)
    table.add_row() # Empty row for spacing
    
    # 2. Preamble
    add_row(table, 'ДОЛУПОДПИСАНИЯТ:', 'НИЖЕПОДПИСАВШИЙСЯ:', bold=True)
    
    # 3. Client details
    add_row(table, 
            'КРАСАВИН ИГОР ПЕТРОВИЧ (KRASAVIN IGOR), гражданин на Руската Федерация, роден на 25.12.1967 г. в Красноярски край (СССР), притежател на задграничен паспорт серия и номер 76 4436769, издаден от МВР 50055 на 23.04.2021 г., валиден до 23.04.2031 г., с постоянен адрес: Руска Федерация, Московска област, гр. Електрогорск, ул. Калинина, дом 5В.', 
            'КРАСАВИН ИГОР ПЕТРОВИЧ (KRASAVIN IGOR), гражданин РФ, дата рождения 25.12.1967, место рождения Красноярский край (СССР), притежатель заграничного паспорта серия и номер 76 4436769, выданного МВД 50055 дата выдачи 23.04.2021 г., действителен до 23.04.2031 г., зарегистрирован по адресу: РФ, Московская обл., г. Электрогорск, ул. Калинина, д. 5В.')
    
    table.add_row()
    
    # 4. Authorize
    add_row(table, 'С НАСТОЯЩОТО УПЪЛНОМОЩАВАМ:', 'НАСТОЯЩИМ УПОЛНОМОЧИВАЮ:', bold=True)
    
    # 5. Lawyer details
    add_row(table, 
            'АЛЕКСАНДРУ ПОПА (ALEXANDRU POPA), румънски гражданин, с постоянен адрес: гр. Букурещ, ул. Виктория № 15, Румъния, притежател на лична карта (Carte de Identitate) серия RX № 458912, издадена от SPCLEP Bucuresti на 15.06.2021 г., ЕГН / CNP 1850615410023, в качеството си на законен представител на A.S. CONSULTING COMPANY S.R.L.', 
            'АЛЕКСАНДРУ ПОПА (ALEXANDRU POPA), гражданин Румынии, адрес регистрации: г. Бухарест, ул. Виктория № 15, Румыния, владелец удостоверения личности (Carte de Identitate) серия RX № 458912, выданного SPCLEP Bucuresti дата выдачи 15.06.2021 г., личный номер (CNP) 1850615410023, в качестве законного представителя A.S. CONSULTING COMPANY S.R.L.')
            
    table.add_row()
    
    # 6. Purpose
    add_row(table, 'Да ме представлява със следните права:', 'Представлять меня со следующими правами:', bold=True)
    
    # 7. Bullet 1
    add_row(table, 
            '- Да ме представлява пред Министерство на правосъдието на Република България (Дирекция „Българско гражданство“), както и пред всички други компетентни институции.', 
            '- Представлять меня в Министерстве юстиции Республики Болгария (Дирекция «Болгарское гражданство»), а также во всех других компетентных учреждениях.')
            
    # 8. Bullet 2
    add_row(table, 
            '- Да подава заявления, да изисква и получава всякакви документи, удостоверения и официална кореспонденция във връзка с моята преписка за придобиване/възстановяване на българско гражданство.', 
            '- Подавать заявления, запрашивать и получать любые документы, справки и официальную переписку в связи с моим делом о получении/восстановлении болгарского гражданства.')
            
    # 9. Bullet 3
    add_row(table, 
            '- Да изисква и получава официална информация относно статуса и хода на разглеждането на моята преписка.', 
            '- Запрашивать и получать официальную информацию о статусе и ходе рассмотрения моего дела.')
            
    # 10. Bullet 4
    add_row(table, 
            '- Да се подписва от мое име и за моя сметка навсякъде, където това е необходимо за изпълнението на настоящото пълномощно.', 
            '- Подписываться от моего имени и за мой счет везде, где это необходимо для выполнения данной доверенности.')
            
    table.add_row()
    
    # 11. Final clause
    add_row(table, 
            'Настоящото пълномощно е безсрочно и важи до изричното му оттегляне.', 
            'Настоящая доверенность является бессрочной и действительна до ее прямого отзыва.')

    doc.add_paragraph()
    doc.add_paragraph()
    
    # Signature block
    sig_table = doc.add_table(rows=1, cols=2)
    sig_table.autofit = True
    
    cell_left = sig_table.cell(0, 0)
    p_left = cell_left.paragraphs[0]
    p_left.add_run('Дата / Дата:\n').bold = True
    p_left.add_run('_____________________')
    
    cell_right = sig_table.cell(0, 1)
    p_right = cell_right.paragraphs[0]
    p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_right.add_run('Упълномощител / Доверитель:\n').bold = True
    p_right.add_run('КРАСАВИН ИГОР ПЕТРОВИЧ\n')
    p_right.add_run('_____________________')

    doc.save('Procura_Bulgaria_Krasavin_Columns.docx')

if __name__ == "__main__":
    create_procura_columns()
