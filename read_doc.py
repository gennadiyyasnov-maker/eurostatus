import docx

doc_path = r"C:\Users\genna\Desktop\инфа про паспорта\Договор гражданство шаблон.docx"
try:
    doc = docx.Document(doc_path)
    text = []
    # Collect paragraph text
    for p in doc.paragraphs:
        if p.text.strip():
            text.append(p.text)
            
    # Also collect table text (since bilingual contracts are usually in tables!)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join([cell.text.replace("\n", " ") for cell in row.cells])
            text.append(row_text)
            
    with open("contract_text.txt", "w", encoding="utf-8") as f:
        f.write("\n".join(text))
    print("Done")
except Exception as e:
    print("Error:", e)
