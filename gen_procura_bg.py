import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_procura():
    doc = docx.Document()

    # Set styling
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # Title
    title = doc.add_paragraph('ПЪЛНОМОЩНО\n(ДОВЕРЕННОСТЬ)')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(14)
    doc.add_paragraph()

    # Preamble 
    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p1.add_run('ДОЛУПОДПИСАНИЯТ / НИЖЕПОДПИСАВШИЙСЯ:\n')
    run.bold = True

    # Client details with filled address
    p1.add_run('КРАСАВИН ИГОР ПЕТРОВИЧ (KRASAVIN IGOR), гражданин на Руската Федерация (гражданин РФ), роден на 25.12.1967 г. в Красноярски край / СССР (дата рождения 25.12.1967, место рождения Красноярский край, СССР), притежател на задграничен паспорт серия и номер 76 4436769, издаден от МВР 50055 на 23.04.2021 г., валиден до 23.04.2031 г., с постоянен адрес (зарегистрирован по адресу): Руска Федерация, Московска област, гр. Електрогорск, ул. Калинина, дом 5В (РФ, Московская обл., г. Электрогорск, ул. Калинина, д. 5В).')
    
    doc.add_paragraph()
    
    # Prin prezenta împuternicesc
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p2.add_run('С НАСТОЯЩОТО УПЪЛНОМОЩАВАМ / НАСТОЯЩИМ УПОЛНОМОЧИВАЮ:\n')
    run.bold = True
    
    # Lawyer details
    p3 = doc.add_paragraph('АЛЕКСАНДРУ ПОПА (ALEXANDRU POPA), румънски гражданин (гражданин Румынии), с постоянен адрес (адрес): гр. Букурещ, ул. Виктория № 15, Румъния (г. Бухарест, ул. Виктория № 15, Румыния), притежател на лична карта (Carte de Identitate) серия RX № 458912, издадена от SPCLEP Bucuresti на 15.06.2021 г., ЕГН / CNP (личный номер) 1850615410023, в качеството си на законен представител на A.S. CONSULTING COMPANY S.R.L. (в качестве законного представителя).')
    p3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_paragraph()
    
    # Scopul / Purpose
    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p4.add_run('Да ме представлява със следните права / Представлять меня со следующими правами:')
    run.bold = True
    
    # Bullet points
    doc.add_paragraph('Да ме представлява пред Министерство на правосъдието на Република България (Дирекция „Българско гражданство“), както и пред всички други компетентни институции. / Представлять меня в Министерстве юстиции Республики Болгария (Дирекция «Болгарское гражданство»), а также во всех других компетентных учреждениях.', style='List Bullet')
    doc.add_paragraph('Да подава заявления, да изисква и получава всякакви документи, удостоверения и официална кореспонденция във връзка с моята преписка за придобиване/възстановяване на българско гражданство. / Подавать заявления, запрашивать и получать любые документы, справки и официальную переписку в связи с моим делом о получении/восстановлении болгарского гражданства.', style='List Bullet')
    doc.add_paragraph('Да изисква и получава официална информация относно статуса и хода на разглеждането на моята преписка. / Запрашивать и получать официальную информацию о статусе и ходе рассмотрения моего дела.', style='List Bullet')
    doc.add_paragraph('Да се подписва от мое име и за моя сметка навсякъде, където това е необходимо за изпълнението на настоящото пълномощно. / Подписываться от моего имени и за мой счет везде, где это необходимо для выполнения данной доверенности.', style='List Bullet')
    
    doc.add_paragraph()
    
    # Final clauses
    p5 = doc.add_paragraph('Настоящото пълномощно е безсрочно и важи до изричното му оттегляне. / Настоящая доверенность является бессрочной и действительна до ее прямого отзыва.')
    p5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Signature block
    table = doc.add_table(rows=1, cols=2)
    table.autofit = True
    
    cell_left = table.cell(0, 0)
    p_left = cell_left.paragraphs[0]
    p_left.add_run('Дата / Дата:\n').bold = True
    p_left.add_run('_____________________')
    
    cell_right = table.cell(0, 1)
    p_right = cell_right.paragraphs[0]
    p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_right.add_run('Подпис / Подпись:\n').bold = True
    p_right.add_run('КРАСАВИН ИГОР ПЕТРОВИЧ\n')
    p_right.add_run('_____________________')

    doc.save('Procura_Bulgaria_Krasavin.docx')

if __name__ == "__main__":
    create_procura()
