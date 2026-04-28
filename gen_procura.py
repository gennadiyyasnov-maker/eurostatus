import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_procura():
    doc = docx.Document()

    # Set styling
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # Title
    title = doc.add_paragraph('PROCURĂ SPECIALĂ\n(SPECIAL POWER OF ATTORNEY)')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(14)
    doc.add_paragraph()

    # Preamble / Subsemnatul
    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p1.add_run('Subsemnatul / The undersigned, ')
    run.bold = True

    # Client details
    p1.add_run('KRASAVIN IGOR, cetățean al Federației Ruse, născut la data de 25.12.1967 în Krasnoyarskiy Kray (URSS), posesor al pașaportului rusesc tip P, seria și numărul 76 4436769, eliberat de autoritatea MVD 50055 la data de 23.04.2021, valabil până la data de 23.04.2031, domiciliat în ______________________________________________ (se va completa adresa de domiciliu),')
    
    doc.add_paragraph()
    
    # Prin prezenta împuternicesc
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p2.add_run('Prin prezenta împuternicesc pe / Hereby authorize:')
    run.bold = True
    
    p3 = doc.add_paragraph('Domnul/Doamna [NUMELE AVOCATULUI/REPREZENTANTULUI], cetățean român, domiciliat(ă) în _____________________________________, posesor/posesoare al/a C.I. seria ____ nr. _________, eliberat de _______________ la data de __________, CNP _______________________, în calitate de reprezentant legal din partea A.S. CONSULTING COMPANY S.R.L.')
    p3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_paragraph()
    
    # Scopul / Purpose
    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p4.add_run('Mandatul este acordat pentru următoarele scopuri / The mandate is granted for the following purposes:')
    run.bold = True
    
    # Bullet points
    doc.add_paragraph('Să mă reprezinte cu depline puteri în fața Autorității Naționale pentru Cetățenie (ANC) din cadrul Ministerului Justiției din România, precum și în fața oricăror altor instituții și autorități publice competente.', style='List Bullet')
    doc.add_paragraph('Să depună cereri, să solicite și să ridice orice documente, adeverințe, certificate sau corespondență oficială referitoare la dosarul meu privind cetățenia română.', style='List Bullet')
    doc.add_paragraph('Să solicite și să obțină informații oficiale cu privire la stadiul soluționării dosarului meu de redobândire / acordare a cetățeniei române.', style='List Bullet')
    doc.add_paragraph('Să semneze în numele meu și pentru mine oriunde va fi necesar în legătură cu îndeplinirea acestui mandat, semnătura sa fiindu-mi opozabilă.', style='List Bullet')
    
    doc.add_paragraph()
    
    # Final clauses
    p5 = doc.add_paragraph('Mandatarul meu este autorizat să îndeplinească toate formalitățile legale necesare pentru ducerea la îndeplinire a prezentului mandat, chiar dacă nu sunt menționate expres, mandatul fiind gratuit și netransmisibil.')
    p5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    p6 = doc.add_paragraph('Prezenta procură este valabilă începând de astăzi, data autentificării, și până la revocarea ei expresă.')
    p6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Signature block
    table = doc.add_table(rows=1, cols=2)
    table.autofit = True
    
    cell_left = table.cell(0, 0)
    p_left = cell_left.paragraphs[0]
    p_left.add_run('Data / Date:\n').bold = True
    p_left.add_run('_____________________')
    
    cell_right = table.cell(0, 1)
    p_right = cell_right.paragraphs[0]
    p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_right.add_run('Semnătura / Signature:\n').bold = True
    p_right.add_run('KRASAVIN IGOR\n')
    p_right.add_run('_____________________')

    doc.save('Procura_ANC_Krasavin.docx')

if __name__ == "__main__":
    create_procura()
