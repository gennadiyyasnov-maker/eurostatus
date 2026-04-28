import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_contract():
    doc = docx.Document()
    
    # Set narrow margins to fit dual columns nicely
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Contract de prestări servicii Nr. 204  |  Договор возмездного оказания услуг № 204")
    run.bold = True
    run.font.size = Pt(14)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("București, 28/04/2026  |  Бухарест, 28/04/2026")

    # Table data
    rows_data = [
        (
            "Сetăţean [NUME PRENUME], denumit în continuare “Beneficiarul” pe de o parte, şi A.S. TOP ISO CONSULTING SRL, CUI 33730451, Jud. Argeş, Mun. Piteşti, Bld. Fraţii Goleşti, Nr. 92, reprezentată de directorul Receanu Liliana Cristina, denumit în continuare “Аgent”, care acţionează în baza statutului pe de altă parte, denumiţi împreună “Рărţi”, au încheiat prezentul Contract (denumit în continuare Contract) despre următoarele:",
            "Гражданин [ФИО], именуемый (-ая) в дальнейшем \"Заказчик\", с одной стороны, и A.S. TOP ISO CONSULTING SRL, ИНН 33730451, Jud. Argeş, Mun. Piteşti, Bld. Fraţii Goleşti, Nr. 92, в лице директора Receanu Liliana Cristina, в дальнейшем \"Агент\", действующего на основании устава, с другой стороны, именуемые вместе \"Стороны\", заключили настоящий договор (далее - Договор) о нижеследующем:",
            False
        ),
        ("1. OBIECTUL CONTRACTULUI", "1. ПРЕДМЕТ ДОГОВОРА", True),
        (
            "1.1. “Beneficiarul” încredințează, iar “Agentul” își asumă obligațiile, în numele, în interesul și pe cheltuiala “Beneficiarului”, de a presta servicii juridice „la cheie” pentru asistența în preschimbarea permisului de conducere străin (emis de Federația Rusă) cu un permis de conducere românesc (în continuare „Permisul”), inclusiv programarea la instituțiile abilitate (DRPCIV), asistență în pregătirea dosarului, precum și reprezentarea intereselor până la obținerea documentului final.",
            "1.1. \"Заказчик\" поручает, а \"Агент\" принимает на себя обязательства от имени, в интересах и за счет \"Заказчика\" оказать комплекс юридических услуг «под ключ» по сопровождению процедуры обмена водительского удостоверения Российской Федерации на водительское удостоверение Румынии (далее — «Права»), включая запись в уполномоченные органы (DRPCIV), помощь в подготовке досье, а также представление интересов вплоть до получения итогового документа.",
            False
        ),
        ("2. RECOMPENSA AGENTULUI, ORDINEA DE PLATĂ", "2. АГЕНТСКОЕ ВОЗНАГРАЖДЕНИЕ, ПОРЯДОК ОПЛАТЫ", True),
        (
            "2.1. Remunerarea „Agentului” în baza prezentului Contract este de 400 euro.",
            "2.1. Вознаграждение \"Агента\" по настоящему Договору составляет 400 евро.",
            False
        ),
        (
            "2.2. Plata serviciilor se efectuează în două etape.",
            "2.2. Оплата услуг осуществляется в два этапа.",
            False
        ),
        (
            "2.2.1. „Beneficiarul” achită „Agentului” un avans de 10% în cuantum dе 40 euro în termen de 2 zile de la data semnării prezentului Contract.",
            "2.2.1. \"Заказчик\" оплачивает \"Агенту\" аванс 10% в размере 40 евро в течение 2 дней с момента подписания данного Договора.",
            False
        ),
        (
            "2.2.2. „Beneficiarul” se obligă să achite „Agentului” valoarea rămasă a serviciilor în sumă de 360 euro imediat după depunerea cu succes a dosarului și fotografierea „Beneficiarului” la Direcția Regim Permise de Conducere (DRPCIV) din România.",
            "2.2.2. \"Заказчик\" обязуется оплатить \"Агенту\" оставшуюся стоимость услуг в размере 360 евро сразу после успешной подачи досье и фотографирования \"Заказчика\" в Управлении по выдаче водительских удостоверений (DRPCIV) Румынии.",
            False
        ),
        (
            "2.2.3. Calculul conform prezentului Contract este efectuat în mod convenabil pentru ambele părți (transfer bancar, numerar, criptomonedă etc.), comisioanele băncare fiind plătite de „Beneficiar”.",
            "2.2.3. Расчет по настоящему Договору производится любым удобным для обеих сторон образом (банковский перевод, наличные, криптовалюта и т.д.), комиссии банка оплачиваются \"Заказчиком\".",
            False
        ),
        ("3. OBLIGAŢIILE PĂRŢILOR", "3. ОБЯЗАННОСТИ СТОРОН", True),
        ("3.1. “Agentul” se obligă:", "3.1. \"Агент\" обязуется:", True),
        (
            "- să analizeze actele “Beneficiarului” necesare pentru preschimbarea permisului;\n- să asigure traducerea legalizată și notarizarea permisului de conducere rus;\n- să coordoneze procedura de obținere a fișei medicale auto în România pentru „Beneficiar”;\n- să asigure programarea și însoțirea „Beneficiarului” la DRPCIV pentru depunerea dosarului;",
            "- провести анализ документов \"Заказчика\", необходимых для обмена прав;\n- обеспечить легализованный перевод и нотариальное заверение российского водительского удостоверения;\n- скоординировать процедуру прохождения автомобильной медицинской комиссии в Румынии;\n- обеспечить запись и личное сопровождение \"Заказчика\" в DRPCIV для подачи досье;",
            False
        ),
        ("3.2. “Beneficiarul” se obligă:", "3.2. \"Заказчик\" обязуется:", True),
        (
            "- să ofere “Agentului” permisul original, actul de identitate românesc (buletin) și alte informații necesare;\n- să se prezinte personal la clinica medicală și la DRPCIV în ziua programată;\n- să achite comisionul către “Agent” în conformitate cu condiţiile prezentului Contract.",
            "- предоставить \"Агенту\" оригинал водительского удостоверения РФ, румынское удостоверение личности (булетин) и иную информацию;\n- лично явиться в медицинскую клинику и в DRPCIV в назначенную дату;\n- оплатить вознаграждение \"Агенту\" в соответствии с условиями настоящего Договора.",
            False
        ),
        ("4. RĂSPUNDEREA PĂRŢILOR", "4. ОТВЕТСТВЕННОСТЬ СТОРОН", True),
        (
            "4.1. Părţile răspund pentru neîndeplinirea sau îndeplinirea necorespunzătoare a obligaţiilor asumate prin prezentul Contract, în conformitate cu legislaţia în vigoare la România. Agentul nu este responsabil dacă preschimbarea este refuzată din cauza interdicțiilor rutiere aplicate Beneficiarului în statul emitent.",
            "4.1. За неисполнение или ненадлежащее исполнение своих обязательств по настоящему договору стороны несут ответственность в соответствии с действующим законодательством Румынии. Агент не несет ответственности, если в обмене отказано по причине наличия у Заказчика запретов на управление ТС в стране выдачи.",
            False
        ),
        ("5. ÎMPREJURĂRILE DE FORŢA-MAJORĂ", "5. ФОРС-МАЖОРНЫЕ ОБСТОЯТЕЛЬСТВА", True),
        (
            "5.1. Părţile sunt eliberate de răspundere pentru neîndeplinirea obligaţiilor din cauza forţei majore (război, calamităţi naturale, modificări ale legislației rutiere, sistarea activității DRPCIV etc.).",
            "5.1. Стороны освобождаются от ответственности за неисполнение обязательств из-за форс-мажора (война, стихийные бедствия, изменения в ПДД и законодательстве, приостановка работы DRPCIV и т.д.).",
            False
        ),
        ("6. SOLUŢIONAREA LITIGIILOR", "6. РАЗРЕШЕНИЕ СПОРОВ", True),
        (
            "6.1. Toate divergenţele se soluționează pe cale amiabilă. În caz contrar, se transmit instanțelor din România.",
            "6.1. Все разногласия решаются путем переговоров. В противном случае спор передается в суды Румынии.",
            False
        ),
        ("7. CONDIŢII SPECIALE", "7. ОСОБЫЕ УСЛОВИЯ", True),
        (
            "7.1. Contractul intră în vigoare la semnare.\n7.2. Informațiile sunt confidențiale.\n7.3. „Agentul” poate implica terțe persoane (traducători, notari, clinici medicale).",
            "7.1. Договор вступает в силу при подписании.\n7.2. Информация является конфиденциальной.\n7.3. \"Агент\" вправе привлекать третьих лиц (переводчиков, нотариусов, клиники).",
            False
        ),
        (
            "7.4. Garanția returnării fondurilor. În cazul unui refuz nemotivat din partea autorităților române de a preschimba permisul (excluzând cazurile de documente false, antecedente rutiere ascunse), „Agentul” restituie 100% din fondurile achitate de „Beneficiar”.",
            "7.4. Гарантия возврата средств. В случае немотивированного отказа румынских властей в обмене прав (за исключением случаев подделки документов, скрытых лишений прав в РФ), \"Агент\" возвращает 100% уплаченных \"Заказчиком\" средств.",
            False
        ),
        ("8. ADRESA, DATELE ŞI SEMNĂTURILE REPREZENTANŢILOR PĂRŢILOR:", "8. АДРЕСА, РЕКВИЗИТЫ И ПОДПИСИ ПРЕДСТАВИТЕЛЕЙ СТОРОН:", True),
        (
            "Beneficiarul:\n[Nume Prenume]\nPașaport/CNP: \nEliberat: \nAdresa: \nE-mail: \n\nSemnătura ___________________",
            "Заказчик:\n[ФИО]\nЗагранпаспорт/CNP: \nВыдан: \nАдрес: \nE-mail: \n\nПодпись __________________",
            False
        ),
        (
            "Agentul: A.S. TOP ISO CONSULTING SRL\nCod Unic de Înregistrare: 33730451\nRegistry No. J03/1385/2014\nАdresa: Str. Eliberarii 25 Bl. IV28 Sc. B Et. 1 Ap. 27 Cod 900257\nReceanu Liliana Cristina\nWebsite: https://eurostatus.online\n\nSemnătura ___________________",
            "Агент: A.S. TOP ISO CONSULTING SRL\nНалоговый номер: 33730451\nРегистрационный код: J03/1385/2014\nАдрес: Str. Eliberarii 25 Bl. IV28 Sc. B Et. 1 Ap. 27 Cod 900257\nReceanu Liliana Cristina\nСайт: https://eurostatus.online\n\nПодпись_________________",
            False
        )
    ]

    table = doc.add_table(rows=0, cols=2)
    table.style = 'Table Grid'
    
    for ro_text, ru_text, is_bold in rows_data:
        row = table.add_row()
        
        ro_cell = row.cells[0]
        p_ro = ro_cell.paragraphs[0]
        run_ro = p_ro.add_run(ro_text)
        run_ro.bold = is_bold
        
        ru_cell = row.cells[1]
        p_ru = ru_cell.paragraphs[0]
        run_ru = p_ru.add_run(ru_text)
        run_ru.bold = is_bold
        
        # Add padding/spacing
        p_ro.paragraph_format.space_after = Pt(6)
        p_ru.paragraph_format.space_after = Pt(6)
        if is_bold:
            p_ro.paragraph_format.space_before = Pt(12)
            p_ru.paragraph_format.space_before = Pt(12)

    doc.save(r'C:\Users\genna\Desktop\инфа про паспорта\Договор_Замена_Прав.docx')

if __name__ == "__main__":
    create_contract()
